import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_paragraph(doc, text, style=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    return p

def md_to_docx(md_path, docx_path):
    doc = Document()

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Headings
        if line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)

        # Horizontal rule
        elif line.strip() in ('---', '***', '___'):
            doc.add_paragraph('─' * 60)

        # Table: collect all table lines
        elif line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            # Filter separator rows
            rows = [r for r in table_lines if not re.match(r'^\|[\s\-:|]+\|', r)]
            if rows:
                # Parse cells
                parsed = []
                for r in rows:
                    cells = [c.strip() for c in r.strip('|').split('|')]
                    parsed.append(cells)
                if parsed:
                    ncols = max(len(r) for r in parsed)
                    table = doc.add_table(rows=len(parsed), cols=ncols)
                    table.style = 'Table Grid'
                    for ri, row in enumerate(parsed):
                        for ci, cell_text in enumerate(row):
                            if ci < ncols:
                                cell = table.cell(ri, ci)
                                cell.text = strip_inline(cell_text)
                                if ri == 0:
                                    for run in cell.paragraphs[0].runs:
                                        run.bold = True
            continue

        # Bullet list
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, line[2:])

        # Numbered list
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(style='List Number')
            add_inline(p, re.sub(r'^\d+\. ', '', line))

        # Blank line
        elif line.strip() == '':
            pass

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            add_inline(p, line)

        i += 1

    doc.save(docx_path)

def strip_inline(text):
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def add_inline(para, text):
    # Split on bold, italic, inline code
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = para.add_run(token[1:-1])
            run.italic = True
        elif token.startswith('`') and token.endswith('`'):
            run = para.add_run(token[1:-1])
            run.font.name = 'Courier New'
        else:
            para.add_run(token)


output_dir = r'D:\Sandeep_AITest\Test_Claude\StockAnnoucements\output'

md_files = [f for f in os.listdir(output_dir) if f.endswith('.md')]
for md_file in md_files:
    md_path = os.path.join(output_dir, md_file)
    docx_path = os.path.join(output_dir, md_file.replace('.md', '.docx'))
    print(f"Converting {md_file} -> {os.path.basename(docx_path)} ...", end=' ')
    md_to_docx(md_path, docx_path)
    print("done")

print("All conversions complete.")
